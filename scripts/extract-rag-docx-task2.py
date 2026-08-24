from __future__ import annotations

import argparse
import hashlib
import json
import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


PARSER_VERSION = "rag-docx-task2/1.1.0"
INCOMPLETE_MARKERS = ("不完整", "要改", "要搞")
MULTI_MARKERS = ("两篇", "很多篇", "多篇")
TASK2_INSTRUCTION_RE = re.compile(
    r"\b(?:to what extent|do (?:the )?(?:advantages|disadvantages)|do you agree|"
    r"discuss both|"
    r"what (?:are|is|problems|causes|effects|measures|solutions)|"
    r"is (?:this|it) (?:a )?positive|is (?:this|it) (?:a )?negative|"
    r"positive or negative development|advantages outweigh|outweigh the disadvantages)\b",
    flags=re.I,
)
TASK2_WHY_HOW_RE = re.compile(
    r"(?:^|[.!?？]\s+|\n\s*)(?:why|how)\b(?=[^?？\n]{0,160}[?？])",
    flags=re.I,
)


def contains_task2_instruction(text: str) -> bool:
    return bool(TASK2_INSTRUCTION_RE.search(text) or TASK2_WHY_HOW_RE.search(text))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Classify and normalize Task 2 DOCX essays without inventing missing content.")
    parser.add_argument("--essay-dir", type=Path, required=True)
    parser.add_argument("--reference-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--question-overrides", type=Path)
    parser.add_argument("--summary-only", action="store_true")
    return parser.parse_args()


def clean(value: Any) -> str | None:
    if value is None:
        return None
    text = unicodedata.normalize("NFKC", str(value).replace("\xa0", " ").replace("\r", "\n"))
    text = "\n".join(re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text or None


def english_ratio(text: str) -> float:
    latin = len(re.findall(r"[A-Za-z]", text))
    cjk = len(re.findall(r"[\u3400-\u9fff]", text))
    return latin / max(latin + cjk, 1)


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z]+(?:['’-][A-Za-z]+)?", text))


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(re.sub(r"\s+", " ", text).strip().lower().encode("utf-8")).hexdigest()


def extract_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = [clean(paragraph.text) for paragraph in doc.paragraphs]
    paragraphs = [text for text in paragraphs if text]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                rows.append(values)
        if rows:
            tables.append(rows)
    return {"paragraphs": paragraphs, "tables": tables}


def detect_task(file_name: str, text: str) -> str | None:
    sample = f"{file_name}\n{text[:6000]}".lower()
    if "小作文" in file_name:
        return "TASK1"
    if "大作文" in file_name or contains_task2_instruction(sample):
        return "TASK2"
    opening = text[:1600].lower()
    if re.search(r"\b(the (?:bar |line |pie )?(?:chart|graph)|the table|the diagram|the maps?)\b", opening):
        return "TASK1"
    return None


def detect_subtype(question: str | None) -> str | None:
    sample = (question or "").lower()
    if "discuss both" in sample:
        return "DISCUSSION"
    question_parts = len(re.findall(r"[?？]", sample))
    if question_parts >= 2 and re.search(r"\b(?:why|how|what|who)\b", sample):
        return "REPORT_TWO_QUESTIONS"
    if (
        "advantages outweigh" in sample
        or "outweigh the disadvantages" in sample
        or "outweigh" in sample
        or "positive or negative development" in sample
        or "advantages and disadvantages" in sample
        or "benefits of" in sample and "drawbacks" in sample
    ):
        return "ADVANTAGES_DISADVANTAGES"
    if "to what extent" in sample or "how far do you agree" in sample or "agree or disagree" in sample:
        return "OPINION"
    if re.search(r"\b(?:why|how|what|who)\b", sample):
        return "REPORT"
    return None


def strip_question_label(line: str) -> str:
    return re.sub(r"^(?:题目|question)\s*[:：]\s*", "", line, flags=re.I).strip()


def detect_question(lines: list[str]) -> str | None:
    """Join a prompt statement with adjacent Why/How/Outweigh instruction lines."""
    candidates: list[tuple[tuple[int, int, int, int], str]] = []
    opening = [line for line in lines[:24] if line]
    for index, original in enumerate(opening):
        line = strip_question_label(original)
        explicit = line != original
        instruction = contains_task2_instruction(line)
        if not explicit and not instruction:
            continue

        parts = [line]
        # A short instruction such as "Why? How can this be solved?" belongs
        # to the statement immediately before it. Include at most two short,
        # English prompt lines and stop before essay-length prose.
        start_index = index
        if instruction and index <= 3 and word_count(line) <= 45:
            cursor = index - 1
            while cursor >= 0 and len(parts) < 3:
                previous = strip_question_label(opening[cursor])
                if english_ratio(previous) < 0.80 or not 4 <= word_count(previous) <= 60:
                    break
                if re.match(
                    r"^(?:admittedly|however|to begin with|in conclusion|overall|this essay)\b|\bI (?:believe|agree|disagree|think)\b",
                    previous,
                    flags=re.I,
                ):
                    break
                parts.insert(0, previous)
                start_index = cursor
                if "?" in previous or contains_task2_instruction(previous):
                    break
                cursor -= 1

        # Some documents put the statement on one line and two separate
        # instruction questions (e.g. Why...? / How...?) on following lines.
        cursor = index + 1
        while cursor < len(opening) and len(parts) < 4:
            following = strip_question_label(opening[cursor])
            if not contains_task2_instruction(following) or word_count(following) > 45:
                break
            parts.append(following)
            cursor += 1

        candidate = clean(" ".join(parts))
        if candidate and word_count(candidate) >= 8:
            # Prefer an explicitly labelled or early prompt. This prevents a
            # later essay paragraph containing "advantages" from outranking
            # the actual question at the top of the file.
            rank = (
                0 if explicit else 1,
                start_index,
                -(len(TASK2_INSTRUCTION_RE.findall(candidate)) + len(TASK2_WHY_HOW_RE.findall(candidate))),
                -len(parts),
            )
            candidates.append((rank, candidate[:2000]))
    return min(candidates, key=lambda item: item[0])[1] if candidates else None


def split_sentences(paragraph: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z\"'])", paragraph)
    return [part.strip() for part in parts if part.strip()]


def build_offsets(paragraph_texts: list[str]) -> tuple[str, list[dict[str, Any]], list[dict[str, Any]]]:
    essay_text = "\n\n".join(paragraph_texts)
    paragraphs = []
    sentences = []
    cursor = 0
    sentence_index = 0
    for paragraph_index, text in enumerate(paragraph_texts):
        if paragraph_index:
            cursor += 2
        start = cursor
        paragraphs.append({"index": paragraph_index, "text": text, "startOffset": start, "endOffset": start + len(text)})
        local_cursor = 0
        for sentence in split_sentences(text):
            relative = text.find(sentence, local_cursor)
            if relative < 0:
                continue
            absolute = start + relative
            sentences.append({
                "sid": f"s{sentence_index + 1:02d}",
                "index": sentence_index,
                "paragraphIndex": paragraph_index,
                "text": sentence,
                "startOffset": absolute,
                "endOffset": absolute + len(sentence),
            })
            local_cursor = relative + len(sentence)
            sentence_index += 1
        cursor = start + len(text)
    return essay_text, paragraphs, sentences


def structured_reference_essay(extracted: dict[str, Any]) -> list[str]:
    paragraphs = []
    for table in extracted["tables"]:
        if not table:
            continue
        header = [value or "" for value in table[0]]
        english_col = next((index for index, value in enumerate(header) if "英文范文" in value or "english" in value.lower()), None)
        if english_col is None:
            continue
        for row in table[1:]:
            value = row[english_col] if english_col < len(row) else None
            if value and english_ratio(value) > 0.85 and word_count(value) >= 20:
                paragraphs.append(value)
    return paragraphs


def freeform_essay(extracted: dict[str, Any], question: str | None) -> list[str]:
    paragraphs = []
    for text in extracted["paragraphs"]:
        if question and (text == question or (word_count(text) >= 4 and re.sub(r"\s+", " ", text) in re.sub(r"\s+", " ", question))):
            continue
        if re.match(r"^(?:题目|中文题目|整理说明|IELTS TASK 2)", text, flags=re.I):
            continue
        if english_ratio(text) >= 0.85 and word_count(text) >= 20:
            paragraphs.append(text)
    return paragraphs


def parse_document(path: Path, collection: str, question_overrides: dict[str, Any]) -> dict[str, Any]:
    raw_bytes = path.read_bytes()
    extracted = extract_docx(path)
    all_lines = extracted["paragraphs"] + [value for table in extracted["tables"] for row in table for value in row if value]
    raw_text = "\n\n".join(all_lines)
    question = detect_question(extracted["paragraphs"])
    if not question:
        question = detect_question(all_lines[:40])
    question_source = "DOCX_EXTRACTED" if question else None
    question_source_key = None
    override = question_overrides.get(path.name)
    if override:
        question = clean(override.get("questionText"))
        question_source = override.get("sourceType", "CURATED_OVERRIDE")
        question_source_key = override.get("sourceKey")
    task = detect_task(path.name, raw_text)
    if override and question and task is None:
        task = "TASK2"
    paragraph_texts = structured_reference_essay(extracted) if collection == "REFERENCE_MODEL" else freeform_essay(extracted, question)
    essay_text, paragraphs, sentences = build_offsets(paragraph_texts)
    words = word_count(essay_text)
    explicit_incomplete = any(marker in path.name for marker in INCOMPLETE_MARKERS)
    explicit_multi = any(marker in path.name for marker in MULTI_MARKERS)
    conclusion_count = sum(
        bool(re.match(r"^(?:in conclusion|to conclude|overall)\b", paragraph, flags=re.I))
        for paragraph in paragraph_texts
    )
    multi_suspected = explicit_multi or words > 750 or conclusion_count >= 2

    warnings = []
    if task != "TASK2":
        warnings.append(f"task={task or 'UNKNOWN'}")
    if not question:
        warnings.append("missing_question")
    if explicit_incomplete:
        warnings.append("explicit_incomplete_marker")
    if multi_suspected:
        warnings.append("multiple_essays_or_versions_suspected")
    if words < 180:
        warnings.append(f"short_essay={words}")
    if len(paragraphs) < 4:
        warnings.append(f"few_paragraphs={len(paragraphs)}")
    if conclusion_count == 0:
        warnings.append("missing_conclusion_marker")

    if task == "TASK1":
        status = "OUT_OF_SCOPE_TASK1"
    elif multi_suspected:
        status = "MULTI_ESSAY_DOCUMENT"
    elif explicit_incomplete or words < 180 or len(paragraphs) < 3:
        status = "PARTIAL"
    elif not question:
        status = "MISSING_PROMPT"
    elif task != "TASK2":
        status = "NEEDS_REVIEW"
    else:
        status = "COMPLETE"

    # The user confirmed that both Word collections are recommended high-score
    # references, not student gold labels and not the only valid viewpoint.
    # Incomplete documents may still be retrieved at language/paragraph scope.
    rag_scopes = []
    if task == "TASK2" and not multi_suspected and paragraphs:
        rag_scopes.extend(["LANGUAGE", "PARAGRAPH_LOGIC"])
        if question:
            rag_scopes.append("TASK_RESPONSE")
        if status == "COMPLETE":
            rag_scopes.append("FULL_ESSAY")
    allowed_for_rag = bool(rag_scopes)
    content_role = "MODEL_ESSAY" if task == "TASK2" else "OUT_OF_SCOPE"

    return {
        "documentKey": sha256_bytes(raw_bytes)[:24],
        "collection": collection,
        "fileName": path.name,
        "sourcePath": str(path.resolve()),
        "fileHash": sha256_bytes(raw_bytes),
        "textHash": sha256_text(essay_text),
        "task": task,
        "subtype": detect_subtype(question),
        "questionText": question,
        "questionSource": question_source,
        "questionSourceKey": question_source_key,
        "essayText": essay_text,
        "rawExtractedText": raw_text,
        "wordCount": words,
        "paragraphs": paragraphs,
        "sentences": sentences,
        "quality": {
            "status": status,
            "contentRole": content_role,
            "warnings": warnings,
            "allowedForRag": allowed_for_rag,
            "excludeFromEval": True,
            "bandClaim": 9.0 if "9分" in path.name else None,
            "bandVerified": False,
            "curationStatus": "USER_CONFIRMED_HIGH_SCORING_RECOMMENDATION" if task == "TASK2" else None,
            "ragScopes": rag_scopes,
        },
        "stats": {
            "paragraphs": len(paragraphs),
            "sentences": len(sentences),
            "tables": len(extracted["tables"]),
            "rawParagraphs": len(extracted["paragraphs"]),
        },
    }


def main() -> None:
    args = parse_args()
    question_overrides = {}
    if args.question_overrides:
        override_payload = json.loads(args.question_overrides.read_text(encoding="utf-8"))
        question_overrides = override_payload.get("documents", override_payload)
    documents = []
    errors = []
    for directory, collection in ((args.essay_dir, "ESSAY_COLLECTION"), (args.reference_dir, "REFERENCE_MODEL")):
        for path in sorted(directory.resolve().glob("*.docx")):
            if path.name.startswith(("~$", ".~")):
                continue
            try:
                documents.append(parse_document(path, collection, question_overrides))
            except Exception as exc:
                errors.append({"sourcePath": str(path.resolve()), "error": f"{type(exc).__name__}: {exc}"})

    seen = {}
    duplicates = []
    for document in documents:
        previous = seen.get(document["textHash"])
        if previous and document["essayText"]:
            duplicates.append({"fileName": document["fileName"], "duplicateOf": previous})
            document["quality"]["warnings"].append(f"duplicate_text_of={previous}")
            document["quality"]["allowedForRag"] = False
            document["quality"]["ragScopes"] = []
        else:
            seen[document["textHash"]] = document["fileName"]

    statuses = {}
    for document in documents:
        status = document["quality"]["status"]
        statuses[status] = statuses.get(status, 0) + 1
    manifest = {
        "schemaVersion": "rag-docx-task2-manifest/1.0.0",
        "parserVersion": PARSER_VERSION,
        "generatedAt": datetime.now().astimezone().isoformat(),
        "summary": {
            "files": len(documents) + len(errors),
            "documents": len(documents),
            "errors": len(errors),
            "task2": sum(document["task"] == "TASK2" for document in documents),
            "allowedForRag": sum(document["quality"]["allowedForRag"] for document in documents),
            "duplicates": len(duplicates),
            "statuses": statuses,
        },
        "qualityQueue": [
            {"fileName": document["fileName"], "collection": document["collection"], "status": document["quality"]["status"], "warnings": document["quality"]["warnings"]}
            for document in documents
            if document["quality"]["status"] != "COMPLETE" or document["quality"]["warnings"]
        ],
        "duplicates": duplicates,
        "errors": errors,
        "documents": documents,
    }
    args.output.resolve().parent.mkdir(parents=True, exist_ok=True)
    args.output.resolve().write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    console = {"summary": manifest["summary"], "errors": errors}
    if not args.summary_only:
        console["qualityQueue"] = manifest["qualityQueue"]
    print(json.dumps(console, ensure_ascii=False, indent=2))
    print(f"Manifest: {args.output.resolve()}")


if __name__ == "__main__":
    main()
