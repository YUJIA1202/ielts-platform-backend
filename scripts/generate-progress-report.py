#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 AI批改RAG系统今日进度报告_2026-06-24.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 样式辅助 ──────────────────────────────────────────────────────────────────
def set_run_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0, 70, 127)):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    set_run_font(run, bold=True, size=(18 if level==1 else 14 if level==2 else 12), color=color)
    return p

def para(text='', bold=False, size=11, indent=0, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, size=size, color=color)
    return p

def bullet(text, level=1, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, size=10.5)
        r2 = p.add_run(text)
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=10, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # blue bg
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0070C0')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        bg = 'DDEEFF' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_run_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(72)
r = p_title.add_run('AI 批改 RAG 系统')
set_run_font(r, bold=True, size=24, color=(0, 70, 127))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('今日进度报告')
set_run_font(r2, bold=True, size=20, color=(0, 112, 192))

doc.add_paragraph()
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_date.add_run('2026-06-24')
set_run_font(r3, size=14, color=(89, 89, 89))

doc.add_paragraph()
p_proj = doc.add_paragraph()
p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_proj.add_run('项目：IELTS Writing AI 批改平台  ·  阶段：V6 数据治理（阶段③）')
set_run_font(r4, size=11, color=(89, 89, 89))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §0  执行摘要
# ══════════════════════════════════════════════════════════════════════════════
heading('§0  执行摘要', 1)
hr()
para('今日完成 V6 方案阶段③（数据治理）的两项核心只读分析：'
     '一是 973 条批注差距的完整重新分类；二是全量 535 篇源文件状态审计。'
     '关键发现：大作文/小作文数字编号范文系列（大作文33‒52、小作文100‒98）'
     '含大量教师蓝色文字，现有 pipeline 完全未捕获，需 RawAnnotationSignal 才能处理。'
     '当前 DB 记录、原始 Word 文件、KnowledgeAnnotation/Chunk/Embedding 均未做任何写操作。',
     size=11)

doc.add_paragraph()
# 数字快览表
heading('关键数字快览', 2)
add_table(
    ['指标', '数值', '说明'],
    [
        ['KnowledgeAnnotation 总量', '12,440', '已全量迁移，幂等，ImportBatch id=3'],
        ['批注差距 (gap)', '973', 'TYPE_A=931 / TYPE_B=42 / TYPE_C=0'],
        ['manifest 文档总数', '535', '含 TEACHER_REVIEW 331 + MODEL_ESSAY 204'],
        ['STALE 文件（源已删）', '3', 'docId=337/433/492，ann=0，待 soft exclude'],
        ['DROP_HIGH_CONFIDENCE', '10', '源存在，全信号为0，待用户确认'],
        ['DROP_MEDIUM_CONFIDENCE', '17', 'MODEL_ESSAY 无信号，需区分范文/草稿'],
        ['DO_NOT_DROP 作业类文件', '77', '含 human override 4 个'],
        ['磁盘未导入文件', '434', '含 33 个作业类，不在 DB，不影响 RAG'],
        ['MODEL_ESSAY 蓝色范文（重大发现）', '~40+', '数字编号系列，0 comments 但有10-138蓝色run'],
    ],
    [4, 2.5, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §1  背景与目标
# ══════════════════════════════════════════════════════════════════════════════
heading('§1  背景与目标', 1)
hr()

heading('1.1  项目背景', 2)
para('本系统为 IELTS 写作 AI 批改平台，核心链路为：教师批改 Word → Python 提取 → manifest.json → '
     'DB 导入 → KnowledgeChunk → Embedding → RAG 检索 → AI 批改输出。'
     'V6 方案目标是将原始教师批注（comment、蓝色文字、行内修改）结构化为'
     'KnowledgeAnnotation，支持 RRF+MMR 检索增强与 30% 分层复核策略。')

heading('1.2  今日目标', 2)
bullet('确认 4 个文件的 human override 分类（ran/may/joan/大作文作业2 z）')
bullet('完成 973 条差距的重新分类（TYPE_A/B/C 细分 + 信号层次估算）')
bullet('完成全量 534 篇源文件状态审计（STALE / DROP / DO_NOT_DROP）')
bullet('建立 20 个人工确认样本（§5 table）')
bullet('生成本份进度报告')

heading('1.3  当前禁止事项（永久约束）', 2)
constraints = [
    '不要 db push',
    '不要新增 RawAnnotationSignal 表',
    '不要新增 KnowledgeEssayUnit 表',
    '不要运行任何 --write',
    '不要 migrate-annotations --full',
    '不要删除原始 Word',
    '不要删除数据库记录',
    '不要修改 KnowledgeAnnotation / KnowledgeChunk / KnowledgeEmbedding',
    '不要修改 task 字段',
    '不要写 document-level quality marks',
    '不要把 sourceType=MODEL_ESSAY 自动当成范文',
    '不要把作业类文件直接删除',
    '不要把同题不同文当重复',
    '不要把蓝色直接等同于错误或建议',
    '不要把黑色直接等同于学生原文',
    '不要把 WORD_COMMENT 直接等同于 ERROR_FEEDBACK 或 STUDENT_ESSAY',
    '不要自动转换 RawAnnotationSignal 到 KnowledgeAnnotation',
]
for c in constraints:
    bullet(c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §2  973 差距重新分类
# ══════════════════════════════════════════════════════════════════════════════
heading('§2  973 条批注差距重新分类', 1)
hr()

heading('2.1  差距类型定义', 2)
add_table(
    ['类型', '定义', '数量', '说明'],
    [
        ['TYPE_A', 'MODEL_ESSAY 从未导入（import 阶段过滤）', '931', 'sourceType=MODEL_ESSAY 的全部 manifest comments'],
        ['TYPE_B', 'import 阶段过滤（极短 feedback/anchor）', '42', '长度 <20 字或纯 anchor 标记'],
        ['TYPE_C', 'migrate 未完成', '0', 'migrate 已全量完成，无残留'],
        ['合计', '—', '973', '—'],
    ],
    [2.5, 6, 2, 4]
)

doc.add_paragraph()
heading('2.2  TYPE_A（931）细分', 2)
add_table(
    ['子类', '数量', '占 TYPE_A %', '典型文件', '处理建议'],
    [
        ['教师批注型（含 human override）', '43', '4.6%', 'ran 双边.docx\nmay 程度同意2.docx', '优先导入，DO_NOT_DROP'],
        ['高价值学生复核型', '65', '7.0%', '大作文作业2 z.docx', 'HIGH_VALUE_STUDENT_REVIEW，待确认后导入'],
        ['学生名 MODEL_ESSAY', '109', '11.7%', 'Goodjoy 系列\nMiriam 系列', '需人工确认是学生作文还是范文'],
        ['其他 MODEL_ESSAY（传统范文）', '774', '83.1%', '大作文33.docx\n小作文100-98系列', 'DO_NOT_DROP，需 RawAnnotationSignal 捕获蓝色内容'],
    ],
    [5, 1.8, 2.5, 4, 3.5]
)

doc.add_paragraph()
heading('2.3  三层信号设计（RawAnnotationSignal）', 2)
para('三层分类架构，每条信号独立分类，不预设语义：')
add_table(
    ['层次', '字段名', '候选值（部分）', '说明'],
    [
        ['WHERE', 'sourceLayer', 'WORD_COMMENT / INLINE_COLORED_TEXT / INLINE_BLUE_TEXT / INLINE_BLACK_TEXT / APPENDED_NOTE / GLOBAL_NOTE / CHUNKTEXT_PARSED / UNKNOWN_SOURCE', '信号来源位置'],
        ['WHAT', 'semanticRole', 'ERROR_FEEDBACK / REWRITE_SUGGESTION / TEACHER_REWRITE / MODEL_TEXT / TEACHING_NOTE / EXAMPLE_EXTENSION / STRUCTURE_COMMENT / GLOBAL_FEEDBACK / UNKNOWN', '信号语义角色'],
        ['WHO', 'targetRole', 'STUDENT_ESSAY / MODEL_ESSAY / TEACHER_REWRITE / PROMPT / PARAGRAPH_SAMPLE / WHOLE_DOCUMENT / UNKNOWN', '信号所针对的内容类型'],
    ],
    [1.5, 3, 7, 4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §3  Human Override 确认记录
# ══════════════════════════════════════════════════════════════════════════════
heading('§3  Human Override 确认记录', 1)
hr()

heading('3.1  已确认 Override（4 个文件）', 2)
add_table(
    ['文件名', 'docId', 'Override 分类', 'semanticRoleHint', '备注'],
    [
        ['ran 双边.docx', '378', 'TEACHING_NOTE  /  DO_NOT_DROP', 'TEACHING_NOTE', '在 manifest，有 65 条 comments'],
        ['may 程度同意2.docx', '377', 'TEACHING_NOTE  /  DO_NOT_DROP', 'TEACHING_NOTE', '在 manifest，有 34 条 comments'],
        ['joan 拓展2.docx', 'N/A', 'TEACHING_NOTE  /  DO_NOT_DROP', 'TEACHING_NOTE', '不在 manifest，属 434 未导入文件，需先导入'],
        ['大作文作业2 z.docx', '434', 'HIGH_VALUE_STUDENT_REVIEW  /  DO_NOT_DROP', 'STUDENT_ESSAY_WITH_REVIEW', 'source 文件存在，0 comments，有其他信号待查'],
    ],
    [4.5, 1.5, 4.5, 4, 4]
)

doc.add_paragraph()
heading('3.2  forceNotAutoEligible 规则', 2)
para('被标记为 TEACHING_NOTE / DO_NOT_DROP 的文件，无论信号计数如何，'
     '均设置 forceNotAutoEligible=true，不纳入 autoEligible 自动评估：')
bullet('原因：教学材料可能含大量范文文字（非学生原文），自动判断易误分类')
bullet('处理路径：需人工审核后，手动决定是否转 KnowledgeAnnotation')
bullet('涉及文件：ran 双边.docx、may 程度同意2.docx、joan 拓展2.docx（待导入）')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §4  源文件状态全量审计
# ══════════════════════════════════════════════════════════════════════════════
heading('§4  源文件状态全量审计结果', 1)
hr()

heading('4.1  审计范围', 2)
bullet('查询 535 篇 manifest 文档，扫描两个源目录的 .docx 文件')
bullet('源目录1：essay collection（TEACHER_REVIEW，764 个 .docx）')
bullet('源目录2：essay collection2（MODEL_ESSAY，204 个 .docx）')
bullet('对全部 134 个零 comment 文件运行 Python DOCX Inspector（检查蓝色/彩色 run）')
bullet('识别 434 个磁盘存在但未导入 manifest 的文件')

doc.add_paragraph()
heading('4.2  STALE 文件（源已删除，DB 残留）', 2)
add_table(
    ['docId', '文件名', 'sourceType', 'chunks', 'ann', 'emb', '建议操作'],
    [
        ['337', '2.18雅思大作文.docx', 'MODEL_ESSAY', '2', '0', '2', 'soft exclude → 确认后 hard_delete_db'],
        ['433', '大作文作业1.docx', 'MODEL_ESSAY', '2', '0', '2', 'soft exclude → 确认后 hard_delete_db'],
        ['492', '小作文作业.docx', 'MODEL_ESSAY', '1', '0', '1', 'soft exclude → 确认后 hard_delete_db'],
    ],
    [1.5, 4.5, 3.5, 1.5, 1.2, 1.2, 5.5]
)
para('注：soft exclude = 设 allowedForRag=false，不删除 DB 记录；'
     'hard_delete_db 需用户二次确认。', size=9.5, color=(127,127,127))

doc.add_paragraph()
heading('4.3  DROP_HIGH_CONFIDENCE（10 个文件）', 2)
para('条件：源文件存在 + docx_comments=0 + blueRuns=0 + coloredRuns=0 + highlightRuns=0 '
     '+ 文件名模式匹配（学生名/作业/日期）+ DB ann=0')
add_table(
    ['docId', '文件名', 'sourceType', '文件名模式'],
    [
        ['110', 'Miriam 大小作文7.22.docx', 'TEACHER_REVIEW', '人名+日期'],
        ['111', 'Miriam大作文+小作文7.14.docx', 'TEACHER_REVIEW', '人名+日期'],
        ['218', '大作文7 (来源2).docx', 'TEACHER_REVIEW', '来源编号'],
        ['226', '小作文4 柱图.docx', 'TEACHER_REVIEW', '学生作文特征'],
        ['430', '大作文7.18.docx', 'MODEL_ESSAY', '日期'],
        ['451', '小作文1.docx', 'MODEL_ESSAY', '序号'],
        ['460', '小作文2 (2).docx', 'MODEL_ESSAY', '序号+重复标记'],
        ['464', '小作文5.docx', 'MODEL_ESSAY', '序号'],
        ['470', '小作文7.18.docx', 'MODEL_ESSAY', '日期'],
        ['471', '小作文72.1.docx', 'MODEL_ESSAY', '序号'],
    ],
    [1.5, 6.5, 4, 5]
)
para('⚠ 以上 10 个文件待用户在 §7 确认表中逐一确认后，再决定是否操作。',
     color=(192, 0, 0), size=10)

doc.add_paragraph()
heading('4.4  DROP_MEDIUM_CONFIDENCE（17 个文件，待确认）', 2)
para('条件：MODEL_ESSAY 无任何信号（comments=0, blue=0），但文件名不含明确学生特征，'
     '可能是空壳范文也可能是无批注学生稿。需人工判断。')
bullet('Goodjoy 小作文新3-15.docx 等 Goodjoy 系列（7 个）')
bullet('优缺点题1.docx、优缺点题2.docx 等优缺点系列（4 个）')
bullet('雅思大作文3.docx、雅思小作文3.docx 等（3 个）')
bullet('其他无特征命名文件（3 个）')

doc.add_paragraph()
heading('4.5  重大发现：MODEL_ESSAY 数字编号范文含大量蓝色内容', 2)

p_warn = doc.add_paragraph()
r_warn = p_warn.add_run('★  关键发现：数字编号范文系列被 pipeline 完全遗漏')
set_run_font(r_warn, bold=True, size=12, color=(192, 0, 0))

para('大作文33‒52（约20个文件）和小作文100‒98系列（约20个文件）：'
     'manifest comment=0，但 Python DOCX Inspector 显示每个文件有 10‒138 个 blue run。'
     '这些是教师用蓝色字体写的范文，当前 Python extraction pipeline 中 node_text() 丢弃了所有 run 样式，'
     '导致蓝色内容被当作普通文本混入 rawText，无法与黑色学生原文区分。')

doc.add_paragraph()
add_table(
    ['文件名', 'manifest comments', 'blue runs', 'colored runs', '判断', '行动'],
    [
        ['大作文33.docx', '0', '66', '66', 'DO_NOT_DROP', '需 RawAnnotationSignal 捕获'],
        ['大作文36(1).docx', '0', '48', '138', 'DO_NOT_DROP', '需 RawAnnotationSignal 捕获'],
        ['大作文38.docx', '0', '70', '70', 'DO_NOT_DROP', '需 RawAnnotationSignal 捕获'],
        ['小作文100-98.docx（代表）', '0', '10-40', '10-40', 'DO_NOT_DROP', '需 RawAnnotationSignal 捕获'],
    ],
    [4.5, 3, 2.2, 2.8, 3, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §5  30% 分层复核策略
# ══════════════════════════════════════════════════════════════════════════════
heading('§5  30% 分层复核策略', 1)
hr()

heading('5.1  策略设计', 2)
para('对 12,440 条 KnowledgeAnnotation 按 reviewPriority 分层复核，目标覆盖率 ≥30%。')
add_table(
    ['优先级', '条件', '数量', '占比', '是否纳入复核', '来源'],
    [
        ['HIGH', 'semanticRole=UNKNOWN 或 sourceLayer=UNKNOWN 或 targetRole=UNKNOWN', '10,040', '80.8%', '100% 纳入', '主因：semanticRole=UNKNOWN 8,959 条'],
        ['MEDIUM', '有信号但置信度 <0.7 （预估）', '~1,500', '~12%', '按比例补足至 ≥30%', '—'],
        ['LOW', '置信度 ≥0.7，低风险', '~900', '~7.2%', '仅最小样本', '—'],
        ['合计（HIGH 已达标）', '—', '≥3,732（30%）', '—', 'HIGH 单独已满足', '—'],
    ],
    [2.5, 6, 2, 2, 3.5, 3]
)

doc.add_paragraph()
heading('5.2  autoEligible 12 条件（同时满足才标记）', 2)
conditions = [
    'comments > 0（有 word comment）',
    'blueRuns = 0（无蓝色文字干扰）',
    'coloredRuns ≤ 3（无大量彩色格式）',
    'contentRole ∈ {REVIEW_EXAMPLE, MODEL_PARAGRAPH}',
    'completenessStatus ∈ {COMPLETE, PARTIAL}',
    'sourceType = TEACHER_REVIEW 或 经人工确认的 MODEL_ESSAY',
    'task ≠ null',
    '文件名不含 作业/assignment/homework',
    'forceNotAutoEligible ≠ true（未被 human override 排除）',
    'DB ann > 0（已有 KnowledgeAnnotation）',
    '无 STALE 标记',
    'reviewPriority ≠ NONE',
]
for i, c in enumerate(conditions, 1):
    bullet(f'条件{i:02d}：{c}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §6  磁盘未导入文件
# ══════════════════════════════════════════════════════════════════════════════
heading('§6  磁盘未导入文件（434 个）', 1)
hr()

heading('6.1  概况', 2)
add_table(
    ['来源目录', '总 .docx 数', '已在 manifest', '未导入', '未导入中作业类'],
    [
        ['essay collection（TEACHER_REVIEW）', '764', '331', '433', '约 25 个'],
        ['essay collection2（MODEL_ESSAY）', '204', '204', '0', '0'],
        ['两目录相同文件名（去重）', '—', '—', '1', '—'],
        ['净计未导入', '—', '—', '434', '33 个'],
    ],
    [5.5, 3, 3, 2.5, 3.5]
)

heading('6.2  处理原则', 2)
bullet('434 个未导入文件不在 DB，不影响当前 RAG 系统运行')
bullet('joan 拓展2.docx 在此列（未入 manifest），用户已确认为教学材料，需单独导入')
bullet('33 个作业类文件：不自动导入，待用户逐一确认后决定是否纳入 RAG')
bullet('其他 401 个：多为 TEACHER_REVIEW，可批量导入，但需先完成 RawAnnotationSignal schema')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §7  待用户确认事项
# ══════════════════════════════════════════════════════════════════════════════
heading('§7  待用户确认事项', 1)
hr()

heading('Q1：DROP_HIGH_CONFIDENCE 10 个文件', 2)
para('以下 10 个文件满足所有 DROP_HIGH_CONFIDENCE 条件。'
     '请确认：这些文件是否确实是无价值的未批注学生草稿？', bold=False)
add_table(
    ['docId', '文件名', 'sourceType', '你的判断'],
    [
        ['110', 'Miriam 大小作文7.22.docx', 'TEACHER_REVIEW', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['111', 'Miriam大作文+小作文7.14.docx', 'TEACHER_REVIEW', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['218', '大作文7 (来源2).docx', 'TEACHER_REVIEW', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['226', '小作文4 柱图.docx', 'TEACHER_REVIEW', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['430', '大作文7.18.docx', 'MODEL_ESSAY', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['451', '小作文1.docx', 'MODEL_ESSAY', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['460', '小作文2 (2).docx', 'MODEL_ESSAY', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['464', '小作文5.docx', 'MODEL_ESSAY', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['470', '小作文7.18.docx', 'MODEL_ESSAY', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
        ['471', '小作文72.1.docx', 'MODEL_ESSAY', '□ 确认 DROP  □ DO_NOT_DROP  □ 需进一步检查'],
    ],
    [1.5, 5, 3.5, 7]
)

doc.add_paragraph()
heading('Q2：STALE 文件 3 个（源已删除）', 2)
para('以下 3 个文件的源 Word 已从磁盘删除，DB/manifest 仍有记录（ann=0）。'
     '建议先 soft exclude（allowedForRag=false），再二次确认后 hard_delete_db。')
add_table(
    ['docId', '文件名', '建议', '你的决定'],
    [
        ['337', '2.18雅思大作文.docx', 'soft exclude 后 hard_delete_db', '□ 同意  □ 仅 soft exclude  □ 暂不操作'],
        ['433', '大作文作业1.docx', 'soft exclude 后 hard_delete_db', '□ 同意  □ 仅 soft exclude  □ 暂不操作'],
        ['492', '小作文作业.docx', 'soft exclude 后 hard_delete_db', '□ 同意  □ 仅 soft exclude  □ 暂不操作'],
    ],
    [1.5, 5, 5, 6]
)

doc.add_paragraph()
heading('Q3：RawAnnotationSignal schema — 何时 db push？', 2)
bullet('当前方案已设计完整，仅等用户批准后执行 prisma db push')
bullet('新增字段（不新增表）：reviewPriority / autoEligible / sampledForReview / reviewStatus / reviewedBy / reviewedAt / reviewNotes / essayUnitId（预留）')
bullet('□ 同意本周执行 db push  □ 继续等待  □ 需要先看 schema diff')

doc.add_paragraph()
heading('Q4：DROP_MEDIUM_CONFIDENCE 17 个文件', 2)
para('这 17 个 MODEL_ESSAY 文件无任何信号（0 comments，0 blue runs）。'
     'Goodjoy 系列可能是学生作文（学生名），优缺点系列可能是教师范文。请确认：')
bullet('□ Goodjoy 系列 → 确认是学生作文，可 DROP')
bullet('□ Goodjoy 系列 → 是教师范文，DO_NOT_DROP')
bullet('□ 优缺点系列 → 教师范文但内容为空，可 DROP')
bullet('□ 其他 → 需我进一步调查每个文件')

doc.add_paragraph()
heading('Q5：joan 拓展2.docx 导入时机', 2)
bullet('该文件确认为教学材料（TEACHING_NOTE / DO_NOT_DROP），当前在 434 未导入文件中')
bullet('□ 本轮导入（RawAnnotationSignal schema push 后）')
bullet('□ 暂不导入，等 human review pipeline 完成')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §8  今日完成脚本清单
# ══════════════════════════════════════════════════════════════════════════════
heading('§8  今日完成脚本清单', 1)
hr()

add_table(
    ['脚本', 'npm 命令', '类型', '核心输出'],
    [
        ['reconcile-source-quality-dry-run.ts',
         'npm run rag:reconcile-quality',
         '只读',
         '§A human override / §B TYPE_A 细分 / §C DROP 候选 / §D 30% 采样 / §E 20样本'],
        ['audit-source-file-status-dry-run.ts',
         'npm run rag:audit-file-status',
         '只读',
         '535文档状态 / STALE 3 / DROP_HIGH 10 / DROP_MEDIUM 17 / 未导入 434'],
        ['audit-docx-annotation-signals-dry-run.ts',
         'npm run rag:audit-signals',
         '只读（已有）',
         '20样本 DOCX 蓝色信号检查，7节输出'],
        ['reconcile-annotation-gap-dry-run.ts',
         'npm run rag:reconcile-gap',
         '只读（已有）',
         'TYPE_A=931 / TYPE_B=42 / TYPE_C=0，差距根因'],
    ],
    [6, 4, 1.8, 6.5]
)

doc.add_paragraph()
heading('8.1  Python DOCX Inspector（内嵌函数）', 2)
para('两个脚本均内嵌同一段 Python 字符串，通过 child_process.execSync 调用：')
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Cm(1)
r_code = p_code.add_run(
    'import sys, json\n'
    'sys.stdout.reconfigure(encoding="utf-8")\n'
    'from zipfile import ZipFile, BadZipFile\n'
    'from xml.etree import ElementTree as ET\n'
    'W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"\n'
    'BLUE = {"0000ff","0070c0","4472c4","2e74b5","4f81bd","1f497d",...}\n'
    '# 输出: {"docxCommentCount":N, "blueRuns":N, "coloredRuns":N, "highlightRuns":N}'
)
r_code.font.name = 'Courier New'
r_code._r.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
r_code.font.size = Pt(9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §9  下一步计划
# ══════════════════════════════════════════════════════════════════════════════
heading('§9  下一步计划（按优先级）', 1)
hr()

add_table(
    ['优先级', '任务', '前置条件', '预计工作量'],
    [
        ['P0', '用户确认 §7 Q1‒Q5（人工判断）', '无', '人工，<1小时'],
        ['P0', 'prisma db push RawAnnotationSignal 字段', '用户批准 Q3', '5分钟'],
        ['P1', 'rag:classify-quality --write（写文档级质量标记）', 'schema push', '30分钟'],
        ['P1', '3 个 STALE 文件 soft exclude', '用户确认 Q2', '10分钟'],
        ['P1', 'DROP_HIGH_CONFIDENCE 10 文件处理', '用户确认 Q1', '20分钟'],
        ['P2', 'Python extraction 增强（捕获蓝色文字）', 'schema push', '2-4小时'],
        ['P2', 'RawAnnotationSignal 批量导入（Phase 1）', '提取增强完成', '3-5小时'],
        ['P3', 'joan 拓展2.docx 单独导入', '用户批准 Q5', '1小时'],
        ['P3', '434 未导入文件批量导入评估', 'P2 完成后', '待评估'],
        ['P4', '阶段④ 数据切分（题型分层）', '治理完成', '1-2天'],
    ],
    [1.5, 6, 4, 3]
)

doc.add_paragraph()
heading('9.1  关键路径', 2)
para('用户确认 §7 → db push → classify-quality --write → '
     'Python 提取增强 → RawAnnotationSignal 导入 → '
     '人工复核 30% 样本 → autoEligible 转 KnowledgeAnnotation（需二次确认）')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §10  14 条核心原则
# ══════════════════════════════════════════════════════════════════════════════
heading('§10  项目核心原则（14 条）', 1)
hr()

principles = [
    ('原则1', 'KnowledgeDocument 是文件容器', 'sourceType 表示文件来源目录，不表示文件实际内容类型。MODEL_ESSAY 目录里可能有学生作文。'),
    ('原则2', 'task=null 字段冻结', '85 篇 task=null 文档不自动修复，根因是 Python detect_task() 仅识别"大作文"全字，但原始文件可能含多篇/混合内容。'),
    ('原则3', 'sourceType ≠ 内容类型', '不能因 sourceType=MODEL_ESSAY 就判断文件是范文；需结合文件名、comments、人工确认。'),
    ('原则4', '文件名含"作业" ≠ DROP', '作业类文件可能含有价值的教师批注（如大作文作业2 z.docx），需人工确认。'),
    ('原则5', '同题不同文 ≠ 重复', '同一题目的不同学生作文不是重复数据，不能删除。'),
    ('原则6', '蓝色 ≠ 错误/建议', '蓝色文字可能是教师范文、教学说明、结构标注等，不能自动判断为 ERROR_FEEDBACK。'),
    ('原则7', '黑色 ≠ 学生原文', '黑色文字可能包含教师行内批注、教学说明段落，需结合上下文判断。'),
    ('原则8', 'WORD_COMMENT ≠ ERROR_FEEDBACK', 'Word comment 可能是结构说明、范文引用、全文评论等，不自动等同于 ERROR_FEEDBACK。'),
    ('原则9', '不自动转换 RawAnnotationSignal', 'RawAnnotationSignal → KnowledgeAnnotation 的转换需人工确认 semanticRole 和 targetRole。'),
    ('原则10', '原始 Word 不删除', '即使文件已完成迁移，原始 Word 源文件不删除（支持重导入）。'),
    ('原则11', 'DB 记录不单方向删除', 'STALE 记录走 soft exclude → 二次确认 → hard_delete_db 流程，不直接删除。'),
    ('原则12', 'KnowledgeChunk 不修改', 'chunk 已生成 embedding，修改 chunk 会使向量失效，禁止修改。'),
    ('原则13', 'migrate-annotations 幂等', '已用 commentId=mv6-{8位hash} 去重，重复运行自动跳过，但不在无 --write 下运行。'),
    ('原则14', 'RawAnnotationSignal schema 待 push', '方案已设计完整，等用户批准后执行 db push，当前仅有只读脚本。'),
]
for code, title, desc in principles:
    p = doc.add_paragraph()
    r1 = p.add_run(f'【{code}】{title}：')
    set_run_font(r1, bold=True, size=10.5, color=(0, 70, 127))
    r2 = p.add_run(desc)
    set_run_font(r2, size=10.5)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §11  附录
# ══════════════════════════════════════════════════════════════════════════════
heading('§11  附录', 1)
hr()

heading('11.1  文件目录结构', 2)
p_dir = doc.add_paragraph()
p_dir.paragraph_format.left_indent = Cm(1)
r_dir = p_dir.add_run(
    'ielts-platform/backend/\n'
    '├── scripts/\n'
    '│   ├── reconcile-source-quality-dry-run.ts     # §2 差距重分类\n'
    '│   ├── audit-source-file-status-dry-run.ts     # §4 全量审计\n'
    '│   ├── audit-docx-annotation-signals-dry-run.ts\n'
    '│   ├── reconcile-annotation-gap-dry-run.ts\n'
    '│   ├── classify-document-quality.ts\n'
    '│   └── migrate-annotations-dry-run.ts\n'
    '├── src/services/\n'
    '│   ├── ragRetriever.ts    # MMR (λ=0.5)\n'
    '│   └── hybridRanker.ts   # RRF (k=60)\n'
    '└── prisma/schema.prisma  # V6 schema（RawAnnotationSignal 字段待 push）'
)
r_dir.font.name = 'Courier New'
r_dir._r.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
r_dir.font.size = Pt(9)

doc.add_paragraph()
heading('11.2  数据库状态', 2)
add_table(
    ['表名', '记录数', '状态'],
    [
        ['KnowledgeDocument', '535', '含 3 STALE，10 DROP_HIGH 候选'],
        ['KnowledgeChunk', '~1,800+', '未修改，原始切片保留'],
        ['KnowledgeAnnotation', '12,440', '全量迁移完成，ImportBatch id=3'],
        ['KnowledgeEmbedding', '与 Chunk 对应', '未修改，embedding 完整'],
        ['ImportBatch', '3', 'id=1(100条) / id=3(12,340条)'],
        ['RetrievalGoldStandard', '0', '待阶段⑥建立'],
        ['ExperimentRun', '0', '待阶段⑦建立'],
    ],
    [5.5, 3, 9.5]
)

# ── 页脚 ───────────────────────────────────────────────────────────────────────
doc.add_paragraph()
hr()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(f'生成时间：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}  ·  '
                        'IELTS Writing AI 批改平台  ·  V6 数据治理阶段③  ·  保密文件')
set_run_font(r_f, size=9, color=(127, 127, 127))

# ── 保存 ───────────────────────────────────────────────────────────────────────
out = r'C:\Users\86182\Desktop\ielts-writing\AI批改RAG系统今日进度报告_2026-06-24.docx'
doc.save(out)
print(f'✅ 报告已生成：{out}')
